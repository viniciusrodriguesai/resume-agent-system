from __future__ import annotations

import io
import tempfile
from pathlib import Path

from resume_ai.settings import Settings

from .security import SafeUpload, validate_upload


class DocumentReader:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def read_upload(
        self,
        filename: str,
        content: bytes,
        reported_type: str | None = None,
    ) -> str:
        upload = validate_upload(filename, content, self.settings, reported_type)
        text = self._read(upload)
        text = text.strip()
        if not text:
            raise ValueError("Não foi possível extrair texto do arquivo")
        return text[: self.settings.max_document_chars]

    def _read(self, upload: SafeUpload) -> str:
        if upload.extension == ".txt":
            return upload.content.decode("utf-8")
        if self.settings.docling_enabled:
            docling_text = self._read_docling(upload)
            if docling_text:
                return docling_text
        if upload.extension == ".pdf":
            return self._read_pdf(upload.content)
        if upload.extension == ".docx":
            return self._read_docx(upload.content)
        raise ValueError("Formato não suportado")

    @staticmethod
    def _read_pdf(content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content), strict=False)
        if len(reader.pages) > 40:
            raise ValueError("PDF com mais de 40 páginas não é aceito")
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    @staticmethod
    def _read_docx(content: bytes) -> str:
        from docx import Document

        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(paragraphs)

    @staticmethod
    def _read_docling(upload: SafeUpload) -> str | None:
        try:
            from docling.document_converter import DocumentConverter
        except Exception:
            return None
        path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(suffix=upload.extension, delete=False) as temp:
                temp.write(upload.content)
                path = Path(temp.name)
            result = DocumentConverter().convert(path)
            markdown = result.document.export_to_markdown()
            return markdown if isinstance(markdown, str) else None
        except Exception:
            return None
        finally:
            if path:
                path.unlink(missing_ok=True)
