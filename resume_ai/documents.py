from __future__ import annotations

import io
import tempfile
from pathlib import Path
from typing import Dict, Tuple

from .config import Settings

class DocumentParser:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or Settings()

    def status(self) -> Dict[str, object]:
        docling_available = False
        try:
            import docling  # noqa: F401
            docling_available = True
        except ImportError:
            pass
        return {
            "docling_enabled": self.settings.enable_docling,
            "docling_available": docling_available,
            "fallbacks": ["pypdf", "python-docx", "plain text"],
        }

    def parse_bytes(self, data: bytes, filename: str) -> Tuple[str, Dict[str, object]]:
        suffix = Path(filename).suffix.lower()
        if self.settings.enable_docling:
            try:
                text = self._parse_docling(data, filename)
                if text.strip():
                    return text, {"engine": "docling", "filename": filename}
            except Exception as exc:
                fallback_warning = str(exc)
            else:
                fallback_warning = ""
        else:
            fallback_warning = "Docling disabled."

        if suffix == ".pdf":
            text = self._parse_pdf(data)
            return text, {
                "engine": "pypdf",
                "filename": filename,
                "warning": fallback_warning,
            }
        if suffix == ".docx":
            text = self._parse_docx(data)
            return text, {
                "engine": "python-docx",
                "filename": filename,
                "warning": fallback_warning,
            }
        if suffix in {".txt", ".md", ".csv"}:
            return data.decode("utf-8", errors="replace"), {
                "engine": "plain-text",
                "filename": filename,
                "warning": fallback_warning,
            }
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    @staticmethod
    def _parse_docling(data: bytes, filename: str) -> str:
        from docling.document_converter import DocumentConverter

        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / filename
            source.write_bytes(data)
            result = DocumentConverter().convert(source)
            return result.document.export_to_markdown()

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        from docx import Document

        document = Document(io.BytesIO(data))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                table_rows.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join([*paragraphs, *table_rows])
